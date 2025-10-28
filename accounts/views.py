import time
import json
import logging
from datetime import datetime
from django.contrib.auth.models import User
from django.http import HttpResponse, JsonResponse
from django.shortcuts import render, redirect, get_object_or_404
from django.contrib import messages
from django.contrib.auth.decorators import login_required
from django.contrib.auth import authenticate, login, logout
from django.db.models import Q, Count
from collections import Counter
import os
import tempfile

# Handle optional imports gracefully
try:
    import fitz  # PyMuPDF
    HAS_PYMUPDF = True
except ImportError:
    HAS_PYMUPDF = False
    fitz = None

try:
    import docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    docx = None

from .forms import UserProfileForm, SignupForm
from .models import *
from .models import UserProfile
from jobs.models import Job, JobCategory, Company
from .advanced_chatbot import advanced_chatbot

# Handle AI analyzer imports gracefully
try:
    from .ai_analyzer import extract_skills, infer_skills_from_text, resume_quality, check_ats_friendliness
    HAS_AI_ANALYZER = True
except ImportError:
    HAS_AI_ANALYZER = False
    # Create dummy functions
    def extract_skills(text): return []
    def infer_skills_from_text(text): return []
    def resume_quality(text): return {"suggestions": [], "readability": {}}
    def check_ats_friendliness(text): return {"has_contact_info": False, "uses_standard_sections": False, "warnings": []}


def _extract_text_from_file(file_path):
    """Extract raw text from a PDF or DOCX file."""
    ext = file_path.rsplit(".", 1)[1].lower()
    text = ""
    try:
        if ext == "pdf" and HAS_PYMUPDF:
            doc = fitz.open(file_path)
            for page in doc:
                text += page.get_text("text") + "\n"
        elif ext == "docx" and HAS_DOCX:
            doc = docx.Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs])
        else:
            # Fallback: return empty string if libraries not available
            return ""
    except Exception as e:
        print(f"Error extracting text from {file_path}: {e}")
        return None
    return text.strip()


# Create your views here.
def signup(request):
    form = SignupForm(request.POST or None)
    if request.method == 'POST':
        if form.is_valid():
            first_name = form.cleaned_data['first_name']
            last_name = form.cleaned_data['last_name']
            username = form.cleaned_data['username']
            email = form.cleaned_data['email']
            role = form.cleaned_data['role']
            dateofbirth = form.cleaned_data['dob']
            password = form.cleaned_data['password1']

            user = User.objects.create_user(
                username=username, 
                password=password, 
                email=email,
                first_name=first_name,
                last_name=last_name
            )
            CustomUser.objects.create(user=user, role=role)
            UserProfile.objects.create(user=user, dateofbirth=dateofbirth)
            messages.success(request, 'Account created successfully. Please log in.')
            return redirect('login')
        else:
            messages.error(request, 'Please correct the errors below.')

    return render(request, 'accounts/signup.html', {'form': form})


def login_view(request):
    if request.user.is_authenticated:
        if request.user.customuser.role == 'company':
            return redirect('company_dashboard')
        else:
            return redirect('dashboard')
    if request.method == 'POST':
        username = request.POST.get('username')
        password = request.POST.get('password')

        if not User.objects.filter(username=username).exists():
            messages.error(request, 'Username does not exist')
            return redirect('login')

        user = authenticate(request, username=username, password=password)
        if user is not None:
            login(request, user)
            temp = CustomUser.objects.get(user=user) if CustomUser.objects.filter(user=user).exists() else None
            if not temp or temp.role == 'applicant':
                return redirect('dashboard')
            if temp.role == 'company':
                if not temp.company:
                    return redirect("profile_completion")
                return redirect("company_dashboard")
        else:
            messages.error(request, 'Invalid username or password')

    return render(request, 'accounts/login.html')


def logout_view(request):
    logout(request)
    return redirect('home')


@login_required
def profile_edit(request):
    """View to edit user profile with AI-powered resume analysis"""
    try:
        profile = request.user.userprofile
    except UserProfile.DoesNotExist:
        profile = UserProfile(user=request.user)

    if request.method == 'POST':
        print(request.POST)
        form = UserProfileForm(request.POST, request.FILES, instance=profile)
        if form.is_valid():
            profile = form.save()
            
            # AI-powered resume analysis if resume was uploaded
            if 'resume' in request.FILES:
                try:
                    from .ai_enhanced import ai_analyzer
                    resume_text = _extract_text_from_file(profile.resume.path)
                    
                    if resume_text:
                        # Extract skills using AI
                        ai_skills = ai_analyzer.extract_skills_from_text(resume_text)
                        
                        # Analyze resume quality
                        quality_analysis = ai_analyzer.analyze_resume_quality(resume_text)
                        
                        # Update profile with AI-extracted skills
                        if ai_skills:
                            existing_skills = [skill.strip() for skill in (profile.skills or '').split(',') if skill.strip()]
                            all_skills = list(set(existing_skills + ai_skills))
                            profile.skills = ', '.join(all_skills)
                            profile.save()
                            
                            messages.success(request, f'Profile updated! AI extracted {len(ai_skills)} skills from your resume.')
                        else:
                            messages.success(request, 'Profile updated! Resume uploaded successfully.')
                        
                        # Store quality analysis for display
                        request.session['resume_analysis'] = {
                            'score': quality_analysis.get('score', 0),
                            'suggestions': quality_analysis.get('suggestions', []),
                            'strengths': quality_analysis.get('strengths', []),
                            'areas_for_improvement': quality_analysis.get('areas_for_improvement', [])
                        }
                    else:
                        messages.warning(request, 'Profile updated, but could not extract text from resume.')
                        
                except Exception as e:
                    print(f"AI resume analysis failed: {e}")
                    messages.success(request, 'Profile updated! (AI analysis temporarily unavailable)')
            else:
                messages.success(request, 'Profile updated successfully!')
            
            return redirect('profile_view')
        else:
            messages.error(request, 'Please correct the errors below.')
    else:
        form = UserProfileForm(instance=profile)

    # Get resume analysis from session if available
    resume_analysis = request.session.pop('resume_analysis', None)

    context = {
        'form': form,
        'profile': profile,
        'resume_analysis': resume_analysis
    }
    return render(request, 'accounts/profile_edit.html', context)


@login_required
def profile_view(request):
    """View to display user profile"""
    try:
        profile = request.user.userprofile
    except UserProfile.DoesNotExist:
        messages.info(request, 'Please complete your profile first.')
        return redirect('profile_edit')

    # Convert skills string to list for better display
    skills_list = []
    if profile.skills:
        skills_list = [skill.strip() for skill in profile.skills.split(',') if skill.strip()]

    context = {
        'profile': profile,
        'skills_list': skills_list,
        'is_own_profile': True  # Since this is the user's own profile view
    }
    return render(request, 'accounts/profile_view.html', context)


def career_advice_view(request):
    """AI-powered career advice view using Gemini"""
    if not request.user.is_authenticated:
        return redirect('login')
    
    # Handle resume upload
    if request.method == 'POST' and 'resume' in request.FILES:
        print("Resume upload detected")
        try:
            profile = request.user.userprofile
            resume_file = request.FILES['resume']
            print(f"Resume file: {resume_file.name}, size: {resume_file.size}")
            
            # Extract text from resume for AI analysis
            resume_text = ""
            try:
                print(f"Attempting to extract text from: {resume_file.name}")
                if resume_file.name.lower().endswith('.pdf'):
                    import PyPDF2
                    # Reset file pointer to beginning
                    resume_file.seek(0)
                    pdf_reader = PyPDF2.PdfReader(resume_file)
                    print(f"PDF has {len(pdf_reader.pages)} pages")
                    for page in pdf_reader.pages:
                        resume_text += page.extract_text()
                    print(f"Extracted {len(resume_text)} characters from PDF")
                elif resume_file.name.lower().endswith(('.doc', '.docx')):
                    import docx
                    # Reset file pointer to beginning
                    resume_file.seek(0)
                    doc = docx.Document(resume_file)
                    print(f"DOCX has {len(doc.paragraphs)} paragraphs")
                    for paragraph in doc.paragraphs:
                        resume_text += paragraph.text + "\n"
                    print(f"Extracted {len(resume_text)} characters from DOCX")
                else:
                    print(f"Unsupported file type: {resume_file.name}")
            except Exception as e:
                print(f"Resume text extraction failed: {e}")
                resume_text = "Resume uploaded but text extraction failed."
            
            # Save resume to profile after extraction
            profile.resume = resume_file
            profile.save()
            
            # Get AI analysis of resume
            try:
                from .gemini_chatbot import gemini_chatbot
                
                resume_analysis_prompt = f"""
                Analyze this resume and provide detailed feedback in JSON format ONLY:

                Resume Text:
                {resume_text[:2000]}

                IMPORTANT: Respond ONLY with valid JSON in this exact format:
                {{
                    "ats_score": 85,
                    "strengths": ["Clear structure", "Relevant experience"],
                    "improvements": ["Add metrics", "Include keywords"],
                    "keywords_missing": ["Python", "Django", "API"],
                    "format_score": 90,
                    "content_score": 80,
                    "recommendations": ["Use bullet points", "Add achievements"]
                }}

                Do not include any text before or after the JSON. Only return the JSON object.
                """
                
                # Convert UserProfile to dictionary format expected by Gemini
                user_profile_dict = {
                    'skills': profile.skills or '',
                    'first_name': profile.first_name or request.user.first_name or 'User',
                    'resume_text': resume_text,
                    'db_context': ''
                }
                
                response = gemini_chatbot.generate_response(resume_analysis_prompt, user_profile_dict)
                print(f"Gemini response type: {response.get('type')}")
                print(f"Gemini response confidence: {response.get('confidence')}")
                print(f"Gemini response length: {len(response.get('response', ''))}")
                if response.get('response'):
                    print(f"Gemini response preview: {response.get('response')[:200]}...")
                
                if response and response.get('response'):
                    import json
                    try:
                        # Clean the response - remove markdown code blocks if present
                        response_text = response['response'].strip()
                        if response_text.startswith('```json'):
                            response_text = response_text[7:]  # Remove ```json
                        if response_text.startswith('```'):
                            response_text = response_text[3:]   # Remove ```
                        if response_text.endswith('```'):
                            response_text = response_text[:-3]  # Remove trailing ```
                        response_text = response_text.strip()
                        print(f"Cleaned resume response: {response_text[:200]}...")
                        
                        # Try to parse as JSON
                        ai_data = json.loads(response_text)
                        
                        # Store analysis in session for display
                        request.session['resume_analysis'] = {
                            'ats_score': ai_data.get('ats_score', 75),
                            'strengths': ai_data.get('strengths', []),
                            'improvements': ai_data.get('improvements', []),
                            'keywords_missing': ai_data.get('keywords_missing', []),
                            'format_score': ai_data.get('format_score', 80),
                            'content_score': ai_data.get('content_score', 75),
                            'recommendations': ai_data.get('recommendations', [])
                        }
                        
                        messages.success(request, 'Resume analyzed successfully! Check the results below.')
                        
                    except json.JSONDecodeError as e:
                        print(f"JSON decode error: {e}")
                        print(f"Raw response: {response['response']}")
                        # Fallback analysis
                        request.session['resume_analysis'] = {
                            'ats_score': 75,
                            'strengths': ['Clear structure', 'Relevant experience', 'Good formatting'],
                            'improvements': ['Add more quantifiable achievements', 'Include relevant keywords', 'Optimize for ATS'],
                            'keywords_missing': ['Technical skills', 'Industry terms', 'Action verbs'],
                            'format_score': 80,
                            'content_score': 75,
                            'recommendations': ['Use bullet points', 'Include metrics', 'Add skills section']
                        }
                        
            except Exception as e:
                print(f"Resume analysis failed: {e}")
                messages.warning(request, 'Resume uploaded but AI analysis failed. Please try again.')
            
            return redirect('career_advice')
            
        except Exception as e:
            print(f"Resume upload failed: {e}")
            messages.error(request, 'Failed to upload resume. Please try again.')
            return redirect('career_advice')
    
    try:
        profile = request.user.userprofile
        user_skills = []
        experience_years = 0
        
        # Get user skills
        if profile.skills:
            user_skills = [skill.strip() for skill in profile.skills.split(',') if skill.strip()]
        
        # Calculate experience years (simplified)
        if profile.dateofbirth:
            from datetime import date
            today = date.today()
            age = today.year - profile.dateofbirth.year
            # Assume work experience starts at 22
            experience_years = max(0, age - 22)
        
        # Initialize AI responses
        career_advice = None
        profile_insights = None
        ats_score = None
        recommendations = None
        resume_analysis = request.session.get('resume_analysis')
        
        # Get AI-powered insights using Gemini
        if user_skills:
            try:
                from .gemini_chatbot import gemini_chatbot
                
                # Create comprehensive prompt for career advice
                career_prompt = f"""
                As a career advisor, analyze this user's profile and provide comprehensive career guidance in JSON format ONLY:

                User Profile:
                - Skills: {', '.join(user_skills)}
                - Experience Level: {experience_years} years
                - Name: {profile.first_name or ''} {profile.last_name or ''}
                - Phone: {profile.phone or 'Not provided'}
                - Email: {profile.email or 'Not provided'}
                - Resume Available: {'Yes' if profile.resume else 'No'}

                IMPORTANT: Respond ONLY with valid JSON in this exact format:
                {{
                    "profile_completeness": 75,
                    "market_position": "mid",
                    "growth_potential": "medium",
                    "strengths": ["Technical skills", "Problem solving"],
                    "recommendations": ["Get certified", "Build portfolio"],
                    "next_steps": ["Update LinkedIn", "Apply to jobs"],
                    "skill_gaps": ["Cloud", "DevOps"],
                    "market_insights": ["High demand", "Remote work"],
                    "ats_score": 80,
                    "recommendations_list": ["Use keywords", "Add metrics"]
                }}

                Do not include any text before or after the JSON. Only return the JSON object.
                Focus on practical, actionable advice for someone with {experience_years} years of experience in {', '.join(user_skills[:3])}.
                """
                
                # Convert UserProfile to dictionary format expected by Gemini
                user_profile_dict = {
                    'skills': profile.skills or '',
                    'first_name': profile.first_name or request.user.first_name or 'User',
                    'resume_text': '',
                    'db_context': ''
                }
                
                response = gemini_chatbot.generate_response(career_prompt, user_profile_dict)
                print(f"Career advice response type: {response.get('type')}")
                print(f"Career advice response confidence: {response.get('confidence')}")
                if response.get('response'):
                    print(f"Career advice response preview: {response.get('response')[:200]}...")
                
                if response and response.get('response'):
                    import json
                    try:
                        # Clean the response - remove markdown code blocks if present
                        response_text = response['response'].strip()
                        if response_text.startswith('```json'):
                            response_text = response_text[7:]  # Remove ```json
                        if response_text.startswith('```'):
                            response_text = response_text[3:]   # Remove ```
                        if response_text.endswith('```'):
                            response_text = response_text[:-3]  # Remove trailing ```
                        response_text = response_text.strip()
                        print(f"Cleaned career response: {response_text[:200]}...")
                        
                        # Try to parse as JSON
                        ai_data = json.loads(response_text)
                        
                        # Structure the data for template
                        profile_insights = {
                            'profile_completeness': ai_data.get('profile_completeness', 75),
                            'market_position': ai_data.get('market_position', 'mid'),
                            'growth_potential': ai_data.get('growth_potential', 'medium'),
                            'strengths': ai_data.get('strengths', [])
                        }
                        
                        career_advice = {
                            'recommendations': ai_data.get('recommendations', []),
                            'next_steps': ai_data.get('next_steps', []),
                            'skill_gaps': ai_data.get('skill_gaps', []),
                            'market_insights': ai_data.get('market_insights', [])
                        }
                        
                        ats_score = ai_data.get('ats_score', 80)
                        recommendations = ai_data.get('recommendations_list', [])
                        
                    except json.JSONDecodeError as e:
                        print(f"Career advice JSON decode error: {e}")
                        print(f"Raw response: {response['response']}")
                        # Fallback if response isn't valid JSON
                        profile_insights = {
                            'profile_completeness': 75,
                            'market_position': 'mid',
                            'growth_potential': 'medium',
                            'strengths': ['Strong technical skills', 'Good problem-solving ability', 'Adaptable to new technologies']
                        }
                        
                        career_advice = {
                            'recommendations': ['Focus on advanced certifications', 'Build a strong portfolio', 'Network with industry professionals'],
                            'next_steps': ['Update your LinkedIn profile', 'Apply to relevant positions', 'Attend industry events'],
                            'skill_gaps': ['Cloud technologies', 'DevOps practices', 'Advanced frameworks'],
                            'market_insights': ['High demand for full-stack developers', 'Remote work opportunities increasing', 'AI/ML skills becoming essential']
                        }
                        
                        ats_score = 80
                        recommendations = ['Optimize resume keywords', 'Include quantifiable achievements', 'Use standard section headers']
                        
            except Exception as e:
                print(f"Gemini career advice failed: {e}")
                # Fallback data
                profile_insights = {
                    'profile_completeness': 70,
                    'market_position': 'mid',
                    'growth_potential': 'medium',
                    'strengths': ['Technical proficiency', 'Problem-solving skills', 'Adaptability']
                }
                
                career_advice = {
                    'recommendations': ['Continue skill development', 'Build professional network', 'Seek mentorship'],
                    'next_steps': ['Complete profile', 'Upload resume', 'Set job alerts'],
                    'skill_gaps': ['Advanced technologies', 'Leadership skills', 'Industry certifications'],
                    'market_insights': ['Growing tech industry', 'Remote work trends', 'Skill-based hiring']
                }
                
                ats_score = 75
                recommendations = ['Use relevant keywords', 'Include metrics', 'Keep format simple']
        
        context = {
            'career_advice': career_advice,
            'profile_insights': profile_insights,
            'ats_score': ats_score,
            'recommendations': recommendations,
            'resume_analysis': resume_analysis,
            'user_skills': user_skills,
            'experience_years': experience_years,
            'has_results': bool(career_advice or profile_insights or resume_analysis),
        }
        
    except UserProfile.DoesNotExist:
        messages.info(request, 'Please complete your profile to get personalized career advice.')
        return redirect('profile_edit')
    
    return render(request, 'accounts/career_advice.html', context)


@login_required
def company_view(request):
    custom_user = getattr(request.user, 'customuser', None)
    company = getattr(custom_user, 'company', None) if custom_user else None

    if not company:
        messages.info(request, 'Please complete your company profile first.')
        return redirect('profile_completion')

    categories = JobCategory.objects.all().order_by('name')
    # Dynamic metrics for dashboard
    active_jobs_count = company.jobs.filter(is_active=True).count()
    closed_jobs_count = company.jobs.filter(is_active=False).count()
    applicants_count = 0
    try:
        from jobs.models import ApplyForJob
        applicants_count = ApplyForJob.objects.filter(job__company=company).count()
    except Exception:
        applicants_count = 0
    # Get company's jobs
    jobs = company.jobs.all().order_by('-created_at')
    
    context = {
        'company': company,
        'jobs': jobs,
        'active_jobs_count': active_jobs_count,
        'closed_jobs_count': closed_jobs_count,
        'applicants_count': applicants_count,
    }
    return render(request, 'Company/company_jobs.html', context)


@login_required
def notifications_view(request):
    """View all notifications"""
    notifications = Notification.objects.filter(user=request.user).order_by('-created_at')
    
    # Mark notifications as read when viewed
    if request.method == 'POST':
        notification_id = request.POST.get('notification_id')
        mark_all_read = request.POST.get('mark_all_read')
        
        if mark_all_read:
            # Mark all notifications as read
            updated_count = notifications.filter(is_read=False).update(is_read=True)
            messages.success(request, f'{updated_count} notifications marked as read.')
        elif notification_id:
            try:
                notification = Notification.objects.get(id=notification_id, user=request.user)
                notification.is_read = True
                notification.save()
                messages.success(request, 'Notification marked as read.')
            except Notification.DoesNotExist:
                messages.error(request, 'Notification not found.')
        return redirect('notifications')
    
    # Pagination
    from django.core.paginator import Paginator
    paginator = Paginator(notifications, 10)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)
    
    context = {
        'notifications': page_obj,
        'total_notifications': notifications.count(),
        'unread_count': notifications.filter(is_read=False).count(),
    }
    
    return render(request, 'accounts/notifications.html', context)


@login_required
def chatbot_view(request):
    """AI Chatbot page"""
    try:
        profile = request.user.userprofile
        user_profile_data = {
            'first_name': profile.first_name,
            'last_name': profile.last_name,
            'email': profile.email,
            'phone': profile.phone,
            'skills': profile.skills,
            'experience': getattr(profile, 'experience', ''),
            'education': getattr(profile, 'education', ''),
            'resume_text': ''
        }
        
        # Get resume text if available
        if profile.resume:
            try:
                resume_text = _extract_text_from_file(profile.resume.path)
                user_profile_data['resume_text'] = resume_text[:1000]  # Limit text length
            except Exception as e:
                print(f"Error extracting resume text: {e}")
        
        # Get suggested questions
        suggested_questions = advanced_chatbot.get_suggested_questions(user_profile_data)
        
        # Get pre-filled question from URL parameter
        initial_question = request.GET.get('q', '')
        
        context = {
            'suggested_questions': suggested_questions,
            'profile': profile,
            'initial_question': initial_question
        }
        
    except UserProfile.DoesNotExist:
        messages.info(request, 'Please complete your profile to use the AI chatbot.')
        return redirect('profile_edit')
    
    return render(request, 'accounts/chatbot.html', context)


@login_required
def chatbot_api(request):
    """API endpoint for chatbot interactions"""
    if request.method != 'POST':
        return JsonResponse({'error': 'Method not allowed'}, status=405)
    
    try:
        data = json.loads(request.body)
        question = data.get('question', '').strip()
        
        if not question:
            return JsonResponse({'error': 'Question is required'}, status=400)
        
        print(f"Chatbot API: Received question: {question}")
        
        # Get user profile data
        try:
            profile = request.user.userprofile
            user_profile_data = {
                'first_name': profile.first_name or '',
                'last_name': profile.last_name or '',
                'email': profile.email or '',
                'phone': profile.phone or '',
                'skills': profile.skills or '',
                'experience': getattr(profile, 'experience', '') or '',
                'education': getattr(profile, 'education', '') or '',
                'resume_text': ''
            }
            print(f"Chatbot API: User profile data: {user_profile_data}")
        except Exception as e:
            print(f"Chatbot API: Error getting user profile: {e}")
            return JsonResponse({'error': 'User profile not found'}, status=400)
        
        # Get resume text if available
        if profile.resume:
            try:
                resume_text = _extract_text_from_file(profile.resume.path)
                user_profile_data['resume_text'] = resume_text[:1000]
                print(f"Chatbot API: Resume text extracted: {len(resume_text)} characters")
            except Exception as e:
                print(f"Chatbot API: Error extracting resume text: {e}")
        
        # Get conversation history from session
        conversation_history = request.session.get('chatbot_history', [])
        print(f"Chatbot API: Conversation history: {len(conversation_history)} entries")

        # Optional: enrich context with live DB data for job intents
        try:
            q_lower = question.lower()
            include_jobs = any(k in q_lower for k in [
                'job', 'jobs', 'opening', 'openings', 'vacancy', 'vacancies', 'hiring', 'apply'
            ])
            db_context_parts = []
            if include_jobs:
                # Simple relevance: filter by user skills appearing in job title or description
                from jobs.models import Job
                skill_terms = []
                if user_profile_data.get('skills'):
                    if isinstance(user_profile_data['skills'], str):
                        skill_terms = [s.strip() for s in user_profile_data['skills'].split(',') if s.strip()]
                    else:
                        skill_terms = user_profile_data['skills']
                jobs_qs = Job.objects.filter(is_active=True).order_by('-created_at')[:50]
                scored = []
                for j in jobs_qs:
                    text = f"{j.title} {getattr(j, 'description', '')} {getattr(j.company, 'name', '')}".lower()
                    score = sum(1 for s in skill_terms if s and s.lower() in text)
                    scored.append((score, j))
                top = [j for score, j in sorted(scored, key=lambda x: x[0], reverse=True)[:5]]
                if not top:
                    top = list(jobs_qs[:5])
                if top:
                    db_context_parts.append("Top jobs:")
                    for j in top:
                        company_name = getattr(j.company, 'name', '') if getattr(j, 'company', None) else ''
                        loc = getattr(j, 'location', '') if hasattr(j, 'location') else ''
                        db_context_parts.append(f"- {j.title} | {company_name} | {loc}")
            if db_context_parts:
                user_profile_data['db_context'] = "\n".join(db_context_parts)
        except Exception as e:
            print(f"Chatbot API: DB enrichment skipped due to error: {e}")
        
        # Generate AI response
        try:
            print("Chatbot API: Generating AI response...")
            print("Chatbot API: Using Gemini-only chatbot...")
            response = advanced_chatbot.generate_response(question, user_profile_data, conversation_history)
            print(f"Chatbot API: AI response generated: {response}")
        except Exception as e:
            print(f"Chatbot API: Error generating AI response: {e}")
            # Fallback response
            response = {
                'response': f"I understand you're asking about '{question}'. Based on your profile, I can help you with career advice, skills analysis, and resume tips. Could you be more specific about what you'd like to know?",
                'confidence': 0.6,
                'type': 'general'
            }
        
        # Store conversation in session
        conversation_entry = {
            'question': question,
            'response': response['response'],
            'timestamp': datetime.now().isoformat(),
            'type': response.get('type', 'general')
        }
        conversation_history.append(conversation_entry)
        
        # Keep only last 10 conversations
        if len(conversation_history) > 10:
            conversation_history = conversation_history[-10:]
        
        request.session['chatbot_history'] = conversation_history
        
        return JsonResponse({
            'response': response['response'],
            'confidence': response.get('confidence', 0.5),
            'type': response.get('type', 'general'),
            'timestamp': conversation_entry['timestamp']
        })
        
    except json.JSONDecodeError as e:
        print(f"Chatbot API: JSON decode error: {e}")
        return JsonResponse({'error': 'Invalid JSON'}, status=400)
    except Exception as e:
        print(f"Chatbot API: Unexpected error: {e}")
        import traceback
        traceback.print_exc()
        return JsonResponse({'error': f'Internal server error: {str(e)}'}, status=500)


@login_required
def job_alerts(request):
    """Manage job alerts"""
    alerts = JobAlert.objects.filter(user=request.user).order_by('-created_at')
    
    if request.method == 'POST':
        action = request.POST.get('action')
        
        if action == 'create':
            keywords = request.POST.get('keywords', '').strip()
            location = request.POST.get('location', '').strip()
            salary_min = request.POST.get('salary_min', '').strip()
            salary_max = request.POST.get('salary_max', '').strip()
            frequency = request.POST.get('frequency', 'weekly')
            
            if keywords:
                alert = JobAlert.objects.create(
                    user=request.user,
                    keywords=keywords,
                    location=location,
                    salary_min=float(salary_min) if salary_min else None,
                    salary_max=float(salary_max) if salary_max else None,
                    frequency=frequency
                )
                messages.success(request, 'Job alert created successfully!')
                return redirect('job_alerts')
            else:
                messages.error(request, 'Please provide at least one keyword.')
        
        elif action == 'toggle':
            alert_id = request.POST.get('alert_id')
            try:
                alert = JobAlert.objects.get(id=alert_id, user=request.user)
                alert.is_active = not alert.is_active
                alert.save()
                messages.success(request, f'Job alert {"activated" if alert.is_active else "deactivated"}.')
            except JobAlert.DoesNotExist:
                messages.error(request, 'Job alert not found.')
        
        elif action == 'delete':
            alert_id = request.POST.get('alert_id')
            try:
                alert = JobAlert.objects.get(id=alert_id, user=request.user)
                alert.delete()
                messages.success(request, 'Job alert deleted successfully.')
            except JobAlert.DoesNotExist:
                messages.error(request, 'Job alert not found.')
    
    return render(request, 'accounts/job_alerts.html', {'alerts': alerts})